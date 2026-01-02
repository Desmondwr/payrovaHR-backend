from django.core.management.base import BaseCommand
from django.core.files.base import ContentFile
from contracts.models import ContractTemplate, Contract
from accounts.models import EmployerProfile
from accounts.database_utils import get_tenant_database_alias
from io import BytesIO

class Command(BaseCommand):
    help = 'Creates a default contract template with the provided text.'

    def handle(self, *args, **options):
        # Loop through employers and create template for each
        employers = EmployerProfile.objects.all()
        for employer in employers:
            tenant_db = get_tenant_database_alias(employer)
            
            # DYNAMIC DB CONFIGURATION
            # Since this is a management command, the middleware hasn't run to set up connections.
            # We must manually ensure the connection exists in settings.DATABASES or connections handler.
            # However, standard Django usually requires setup in settings.DATABASES.
            # But likely our project might have a helper or we should use the existing DB_NAME/USER etc
            # and just mutate the connection settings temporarily or rely on how migrations do it.
            
            # Let's inspect how database_utils.get_tenant_database_alias works - 
            # it returns a string alias. If settings.DATABASES doesn't have it, we must add it.
            
            from django.conf import settings
            from django.db import connections
            
            if tenant_db not in settings.DATABASES:
                 # Construct the config based on default but with correct name
                 default_config = settings.DATABASES['default'].copy()
                 default_config['NAME'] = employer.database_name
                 settings.DATABASES[tenant_db] = default_config
            
            # Check if template already exists
            try:
                if not ContractTemplate.objects.using(tenant_db).filter(
                    employer_id=employer.id, 
                    name="Default Fixed-Term Contract",
                    contract_type='FIXED_TERM'
                ).exists():
                    
                    try:
                        from docx import Document
                        doc = Document()
                        doc.add_heading('EMPLOYMENT CONTRACT', 0)

                        text = """
Prepared for: [Employee.FirstName] [Employee.LastName] [Employee.Company]

Created by: [Sender.FirstName] [Sender.LastName] [Sender.Company]

This Employment Contract (the "Contract" or "Employment Contract") states the terms and conditions that govern the contractual agreement between [Sender.Company] (the "Company") having its principal place of business at, [Sender.StreetAddress], [Sender.City], [Sender.State], [Sender.PostalCode] and [Employee.FirstName] [Employee.LastName] (the "Employee") who agrees to be bound by this Contract.

WHEREAS, the Company is engaged in the description of business; and WHEREAS, the Company desires to employ and retain the Services of the Employee according to the terms and conditions herein.

NOW, THEREFORE, In consideration of the mutual covenants and promises made by the parties hereto, the Company and the Employee (individually, each a "Party" and collectively, the "Parties") agree as follows:

1. TERM
The term of this Employment Contract shall commence on (Start.Date) (the "Start Date"). The Employee agrees and acknowledges that, just as they have the right to terminate their employment with the Company at any time for any reason, the Company has the same right, and may terminate their employment with the Company at any time for any reason. Either Party may terminate said employment with written notice to the other Party.

2. DUTIES
The Company shall employ the Employee as [Employee.Title] (the "Position"). The Employee accepts employment with the Company on the terms and conditions outlined in this Employment Contract and agrees to devote their full time and attention (with reasonable periods of illness excepted) to the performance of their duties under this Contract.

The Employee pledges to operate in line with this Contract and in the business' best interests, which may or may not require utilizing their full range of abilities to carry out all job-related responsibilities. The Parties agreed to follow all guidelines, practices, norms, and requirements that are enforced by the Company in the performance of the functions and obligations of the job. Additionally, when working for the Company, the Employee promises to follow all applicable community, municipal, statewide, and national statutes.

In general, the Employee shall perform all the duties outlined in the job description in Exhibit A attached hereto.

3. COMPENSATION
In consideration for the performance of the duties hereunder, the Employee shall be entitled to compensation as follows:

a. The Company shall pay the Employee an annual salary (the "Annual Salary"), also referred to as "Wages." Initially, the Annual Salary shall be at the rate of (Annual.Salary.In.Words) ($Annual_Salary_Amount) per year. The Annual Salary shall be payable in installments, minus the usual and customary payroll deductions for FICA, federal and state withholding, etc., at the times and in the manner in effect in accordance with the usual and customary payroll policies in effect at Company.

b. The Company shall pay and provide to employee retirement plans, health insurance, disability insurance plan benefits, and other benefits generally in effect for salaried employees of the Company beginning on the Effective Hiring Date in accordance with and on the same terms as are generally in effect for employees of the Company.

c. The Employee shall be allowed paid time off for vacation, holidays, and other employee benefits not described above in accordance with the Company policies in general effect for the Company's salaried employees.

4. CONFIDENTIALITY
The Employee shall not (i) disclose to any third party any details regarding the business of the Company, including, without limitation, the names of any of its customers, the prices it obtains, the prices at which it sells products, its manner of operation, its plans, its strategies, any of the Company's trade secrets or any other information pertaining to the business of the Company (the "Confidential Information"), (ii) make copies of any Confidential Information or any content based on the concepts contained within the Confidential Information for personal use or for distribution unless requested to do so by the Company, or (iii) use Confidential Information other than solely for the benefit of the Company.

Employee acknowledges and agrees that any legal remedy for any breach of this confidentiality provision may be inadequate and, in the event of any such breach, the Company shall be entitled to immediate and permanent injunctive relief to preclude and/or any such breach (in addition to any remedies at law to which the Company may be entitled) without the posting of any bond or security therefore.

5. RETURN OF PROPERTY
Within seven (7) days of the termination of this Contract, whether by expiration or otherwise, the Employee agrees to return to the Company, all products, samples, or models, and all documents, retaining no copies or notes, relating to the Company's business including, but not limited to:

a. Item 1

b. Item 2

c. Item 3

6. NON-COMPETE AND NON-SOLICITATION
a. Employee hereby agrees (the "Non-Competition Agreement") that, upon the termination of Employee's Employment (for whatever reason, whether during the term of this Agreement or after the termination of this Agreement), for a period of (number of years) following the termination of Employment, Employee shall not directly or indirectly (whether as an officer, director, employee, partner, stockholder, creditor or agent, or representative of other persons or entities) engage in the (Company's industry or type of business) business or in any business in which the Company has, as of the date of such termination, engaged (the "Company's business"), in (Country, State), any county contiguous to such county, and in any county or state in which the Company maintains an office (the "Trade Area").

b. Employee also agrees (the "Non-Solicitation Agreement"), that for a period of (number of years) following the termination of Employee's Employment (for whatever reason, whether during the term of this Agreement or after the termination of this Agreement), Employee shall not directly or indirectly (whether as an officer, director, employee, partner, stockholder, creditor or agent, or representative of other persons or entities) contact or solicit, in any manner indirectly or directly, individuals or entities who were at any time during the original or any extended Term clients of the Company for the purpose of providing (type of services Employee) services by Company during the Term or contact or solicit employees of the Company to seek employment with any person or entity except the Company.

c. Employee agrees that (i) any remedy at law for any breach of the Non-Competition Agreement and/or the Non-Solicitation Agreement would be inadequate, (ii) any breach of the Non-Competition Agreement and/or the Non-Solicitation Agreement shall constitute incontrovertible evidence of irreparable injury to the Company, and (iii) the Company shall be entitled to both immediate and permanent injunctive relief without the necessity of establishing or posting any bond therefore to preclude any such breach (in addition to any remedies of law which the Company may be entitled).

7. EXPENSES
The Employee shall not be entitled to reimbursement for any expenses except those that have been previously approved in writing by the Company. Should the Company require travel by the Employee, the Company shall reimburse the Employee for such travel expenses, along with reasonable lodging and meal expenses upon presentation of receipts for such expenses.

[Sender.Company] promises to reimburse or compensate the Employee once the Employee has provided receipts or expense reports for all the fees incurred by the Employee during their work hours. This includes all the money spent by the Employee while executing their responsibilities mentioned in the Contract. Costs include, without constraints, expenses for items such as commuting, accommodation, attending conferences, and similar items.

This is as per the Employee's retention contract. The personal expenses policy initiatives states that the Business must incorporate this. The Employee shall maintain a list of compensated expenses which the Company shall classify and document as additional pay and which the Employer shall treat as deductible in the remuneration structure of expenditures.

Additionally, and if applicable, the Employer shall give the Employee complete access to transportation and provide them with a vehicle that is currently up to date and fully operational. This shall be for the duration of the Implied Time clause under this Contract. Any and all general driving policies that the Company may at any time implement must be followed by the Employee when using this vehicle. For Employees who require a parking spot for such use, the parking space must be chosen, managed, and upgraded in accordance with the Company's business automobile policy.

8. SEVERABILITY
If a court finds any provision of this Employment Contract invalid or unenforceable, the remainder of this Employment Contract shall be interpreted as best as is possible with respect to the original intent of the Parties.

9. EMPLOYEE REPRESENTATIONS AND WARRANTIES
The Employee represents and warrants to the Company the following:

a. There is no other employment contract or any other contractual obligation to which the Employee is subject, which prevents the Employee from entering into this Contract or from performing fully the Employee's duties under this Contract.

b. The acceptance and submission of this Agreement and the performance of its provisions will (i) not violate any memorandum of understanding or any other documentation to which they are a participant or by which they will be obligated; and (ii) need not be subject to the approval of any individual or persons. Further, the Employee signifies, subpoenas, and commits that (i) their involvement with the Company is in good faith; (ii) their involvement with the Company is for the full term of his authorized employment with the Company; and (iii) their involvement with any authorized private entity as during time frame of their participation with the Company.

c. The Company shall make no specific accommodations for the Employee to perform their duties and responsibilities, other than those specifically described under this Contract.

10. NO MODIFICATION UNLESS IN WRITING
No modification of this Employment Contract shall be valid unless in writing and agreed upon by both Parties.

No aspect of this Contract may be changed, repealed, or dismissed unless both the Employee and the Company consent to it in writing and sign accordingly. No waiver by either Party hereto at any time of any breach or compliance by the other Party hereto with any circumstance or requirement of this Contract to be executed shall be permitted a waiver of similar or identical clauses or constraints for the same or at any preceding or immediately following period. With regard to the assessment hereof, neither Party has made any agreements or statements, whether oral or written, express or implied, that are not stated explicitly in this Agreement.

11. ENTIRE AGREEMENT
This Employment Contract expresses the complete understanding of the Parties with respect to the subject matter and supersedes all prior proposals, agreements, representations, and understandings. This Employment Contract may not be amended except in writing and signed by both Parties.

12. APPLICABLE LAW
This Employment Contract and the interpretation of its terms shall be governed by and construed in accordance with the laws of the State of [Sender.State] and subject to the exclusive jurisdiction of the federal and state courts located in [Sender.Country], [Sender.State].

13. TERMINATION OF THE AGREEMENT
Regardless of any clauses or statements to the contrary in this agreement, the Company or the Employee may end the agreement if they decide — this decision may be made at their discretion and will not result in any remedies, except the ones mentioned in the Contract. Any such termination must be carried out by delivering a formal termination notice to the other Party. The formal notice may be provided through email, written notice, fax, regular mail, or other courier. This notice has to be delivered to the home, fax, or email address provided by the Company or the Employee.

14. ACCEPTANCE FORM
IN WITNESS WHEREOF, each of the Parties has executed this Contract, both Parties by its duly authorized officer, as of the day and year set forth below.

[Sender.FirstName] [Sender.LastName] [Sender.Company] Signature: ____________________ Date: MM / DD / YYYY

[Employee.FirstName] [Employee.LastName] Signature: ____________________ Date: MM / DD / YYYY
                        """
                        
                        # Split text into paragraphs
                        for para in text.split('\\n\\n'):
                            doc.add_paragraph(para.strip())

                        # Save to BytesIO
                        output = BytesIO()
                        doc.save(output)
                        content = ContentFile(output.getvalue())
                    except ImportError:
                         self.stdout.write(self.style.WARNING("python-docx missing. Creating dummy content."))
                         content = ContentFile(b"Dummy DOCX content for template.")
                    
                    template = ContractTemplate(
                        employer_id=employer.id,
                        name="Default Fixed-Term Contract",
                        contract_type="FIXED_TERM",
                        is_default=True
                    )
                    template.file.save("default_fixed_term.docx", content, save=False)
                    template.save(using=tenant_db)
                    self.stdout.write(self.style.SUCCESS(f"Created default template for {employer.company_name}"))
                else:
                    self.stdout.write(f"Template already exists for {employer.company_name}")
            except Exception as e:
                self.stdout.write(self.style.ERROR(f"Error processing {employer.company_name}: {e}"))

        self.stdout.write(self.style.SUCCESS('Successfully seeded default templates.'))

        self.stdout.write(self.style.SUCCESS('Successfully seeded default templates.'))
