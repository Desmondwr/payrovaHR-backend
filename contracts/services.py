import os
from io import BytesIO

from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.utils import timezone

from accounts.models import EmployerProfile
from .models import Contract, ContractTemplate, ContractDocument


def _basic_pdf_bytes(title: str, body: str) -> bytes:
    """Generate a minimal, standards-compliant PDF using only the stdlib (no third-party deps)."""

    def esc(text: str) -> str:
        return (
            text.replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
            .replace("\r", "")
            .replace("\n", "")
        )

    lines = [title, ""] + body.split("\n")
    y = 770  # start near top of letter-sized page

    ops = []
    first = True
    for line in lines:
        if y < 60:
            break
        font_size = 14 if first else 11
        ops.append(f"BT /F1 {font_size} Tf 50 {y} Td ({esc(line)}) Tj ET")
        y -= 18
        first = False

    content = "\n".join(ops).encode("utf-8")

    pdf = bytearray()
    pdf.extend(b"%PDF-1.4\n")
    offsets = []

    def add(obj_bytes: bytes):
        offsets.append(len(pdf))
        pdf.extend(obj_bytes)
        if not pdf.endswith(b"\n"):
            pdf.extend(b"\n")

    add(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    add(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    add(
        b"3 0 obj\n"
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
    )
    add(
        b"4 0 obj\n<< /Length "
        + str(len(content)).encode("ascii")
        + b" >>\nstream\n"
        + content
        + b"\nendstream\nendobj\n"
    )
    add(b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n")

    xref_start = len(pdf)
    pdf.extend(f"xref\n0 {len(offsets)+1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for off in offsets:
        pdf.extend(f"{off:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        b"trailer\n<< /Size "
        + str(len(offsets) + 1).encode("ascii")
        + b" /Root 1 0 R >>\n"
    )
    pdf.extend(b"startxref\n" + str(xref_start).encode("ascii") + b"\n%%EOF\n")
    return bytes(pdf)


def _clean_body(text: str) -> str:
    """
    Remove duplicate heading lines so only one title remains.
    Keeps the first heading that looks like a contract title, drops subsequent ones.
    """
    lines = text.splitlines()
    cleaned = []
    seen_heading = False

    for line in lines:
        stripped = line.strip()
        is_heading = stripped and stripped.upper() == stripped and "CONTRACT" in stripped.upper()

        if is_heading:
            if seen_heading:
                continue  # skip duplicate headings
            seen_heading = True

        cleaned.append(line)

    return "\n".join(cleaned).strip()


def _should_skip_body_line(line: str, employer_info=None, employee_info=None) -> bool:
    """Skip legacy header/address lines from templates to avoid duplication with custom headers."""
    lower = line.strip().lower()
    if not lower:
        return False

    if (
        lower.startswith("created by")
        or lower.startswith("prepared by")
        or lower.startswith("prepared for")
        or lower.startswith("contract parties")
        or lower.startswith("employer date")
        or lower.startswith("employee date")
    ):
        return True

    # Skip placeholder tokens
    if "{co_address" in lower or "{emp_address" in lower or "{co_city" in lower or "{emp_city" in lower:
        return True

    # Skip lines containing actual employer/employee name/address/date values
    for info in (employer_info or {}, employee_info or {}):
        if isinstance(info, dict):
            for key in ("name", "address", "date"):
                val = (info.get(key) or "").strip().lower()
                if not val:
                    continue
                if key == "date":
                    if lower == val:
                        return True
                elif val in lower:
                    return True

    return False


def generate_contract_pdf(contract, template=None):
    """
    Generates a PDF contract document based on a template.

    Args:
        contract (Contract): The contract instance.
        template (ContractTemplate, optional): Specific template to use.
            If None, finds default for contract type.

    Returns:
        ContractDocument: The generated document object.

    Raises:
        ValueError: If no template found.
        ImportError: If python-docx is not installed.
    """

    # -------------------------------------------------------------------------
    # 1) Resolve Template
    # -------------------------------------------------------------------------
    if not template:
        db_alias = contract._state.db or "default"
        template = (
            ContractTemplate.objects.using(db_alias)
            .filter(
                employer_id=contract.employer_id,
                contract_type=contract.contract_type,
                is_default=True,
            )
            .order_by("-created_at")
            .first()
        )
        if not template:
            raise ValueError(f"No default template found for contract type {contract.contract_type}")

    # -------------------------------------------------------------------------
    # 2) Prepare Context (placeholder-style strings for templating)
    # -------------------------------------------------------------------------
    employer_profile = EmployerProfile.objects.using("default").get(id=contract.employer_id)
    User = get_user_model()

    def _get_user_signature_path(user_id):
        """Resolve signature file path + user object from default DB."""
        if not user_id:
            return None, None
        try:
            user_obj = User.objects.using("default").get(id=user_id)
            if user_obj.signature and user_obj.signature.name:
                return user_obj.signature.path, user_obj
        except Exception:
            return None, None
        return None, None

    employer_sig_path, employer_user = _get_user_signature_path(getattr(employer_profile, "user_id", None))

    try:
        employee_obj = contract.employee
    except Exception:
        employee_obj = None

    if employee_obj and getattr(employee_obj, "user_id", None):
        employee_sig_path, employee_user = _get_user_signature_path(employee_obj.user_id)
    else:
        employee_sig_path, employee_user = None, None

    try:
        branch_obj = contract.branch
    except Exception:
        branch_obj = None

    try:
        department_obj = contract.department
    except Exception:
        department_obj = None

    allowances_list = [f"{a.name}: {a.amount} ({a.type})" for a in contract.allowances.all()]
    deductions_list = [f"{d.name}: {d.amount} ({d.type})" for d in contract.deductions.all()]
    allowances_data = [
        {
            "name": getattr(a, "name", ""),
            "amount": getattr(a, "amount", ""),
            "type": getattr(a, "type", ""),
        }
        
        for a in contract.allowances.all()
    ]
    deductions_data = [
        {
            "name": getattr(d, "name", ""),
            "amount": getattr(d, "amount", ""),
            "type": getattr(d, "type", ""),
        }
        for d in contract.deductions.all()
    ]

    context = {
        "{CO_NAME}": employer_profile.company_name,
        "{CO_ADDRESS}": getattr(employer_profile, "physical_address", "") or "",
        "{CO_CITY}": getattr(employer_profile, "company_location", "") or "",
        "{CO_COUNTRY}": getattr(employer_profile, "country", "") or "",
        "{EMP_FIRST}": getattr(employee_obj, "first_name", "") or "",
        "{EMP_LAST}": getattr(employee_obj, "last_name", "") or "",
        "{EMP_TITLE}": getattr(employee_obj, "job_title", "") or "Employee",
        "{EMP_ADDRESS}": getattr(employee_obj, "address", "") or "",
        "{EMP_CITY}": getattr(employee_obj, "city", "") or "",
        "{EMP_COUNTRY}": getattr(employee_obj, "country", "") or "",
        "{BRANCH_NAME}": getattr(branch_obj, "name", "") or "",
        "{DEPARTMENT_NAME}": getattr(department_obj, "name", "") or "",
        "{START_DATE}": contract.start_date.strftime("%Y-%m-%d") if contract.start_date else "",
        "{END_DATE}": contract.end_date.strftime("%Y-%m-%d") if contract.end_date else "",
        "{BASE_SALARY}": f"{contract.base_salary}",
        "{CURRENCY}": contract.currency,
        "{PAY_FREQUENCY}": contract.pay_frequency,
        "{ALLOWANCES}": "; ".join(allowances_list) if allowances_list else "None",
        "{DEDUCTIONS}": "; ".join(deductions_list) if deductions_list else "None",
        "{PROBATION_PERIOD}": getattr(contract, "probation_period", "") if hasattr(contract, "probation_period") else "",
        "{NOTICE_PERIOD}": getattr(contract, "notice_period", "") if hasattr(contract, "notice_period") else "",
        "{HOURS_PER_WEEK}": getattr(contract, "hours_per_week", "") if hasattr(contract, "hours_per_week") else "",
        "{SIGN_DATE_CO}": timezone.now().strftime("%Y-%m-%d"),
        "{SIGN_DATE_EMP}": timezone.now().strftime("%Y-%m-%d"),
    }

    employer_info = {
        "name": employer_profile.company_name,
        "email": getattr(employer_user, "email", "") or "",
        "phone": getattr(employer_profile, "phone_number", "") or "",
        "address": getattr(employer_profile, "physical_address", "") or "",
        "date": timezone.now().strftime("%Y-%m-%d"),
        "sig_path": employer_sig_path,
    }

    employee_info = {
        "name": f"{getattr(employee_obj, 'first_name', '')} {getattr(employee_obj, 'last_name', '')}".strip(),
        "email": getattr(employee_user, "email", "") or getattr(employee_obj, "email", "") or "",
        "phone": getattr(employee_obj, "phone_number", "") or getattr(employee_obj, "phone", "") or "",
        "address": getattr(employee_obj, "address", "") or "",
        "date": timezone.now().strftime("%Y-%m-%d"),
        "sig_path": employee_sig_path,
    }

    signatures_ctx = {
        "employer": {"path": employer_sig_path, "name": employer_info["name"]},
        "employee": {"path": employee_sig_path, "name": employee_info["name"]},
    }

    # -------------------------------------------------------------------------
    # 2b) PDF Renderer (Sweet design, same function/signature)
    # -------------------------------------------------------------------------
    def render_pdf_from_body(title: str, body: str) -> ContentFile:
        """Render a polished PDF using reportlab with placeholder substitution."""
        body = _clean_body(body or "")

        # Replace placeholders
        for key, val in context.items():
            body = body.replace(key, str(val))
            title = title.replace(key, str(val))

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib import colors
        except ImportError:
            # Fallback to a minimal PDF built with stdlib to avoid broken files
            missing = []
            if not signatures_ctx["employer"]["path"]:
                missing.append("Employer signature missing")
            if not signatures_ctx["employee"]["path"]:
                missing.append("Employee signature missing")
            footer = "\n\n".join(missing) if missing else "\n\nSignatures on file."
            return ContentFile(_basic_pdf_bytes(title, body + "\n\n" + footer))

        def wrap_lines(text, font_name, font_size, max_width, canv):
            """Wrap text to fit within max_width using the canvas width metrics."""
            words = text.split()
            if not words:
                return [""]

            lines = []
            current = ""
            for word in words:
                candidate = f"{current} {word}".strip()
                if canv.stringWidth(candidate, font_name, font_size) <= max_width:
                    current = candidate
                else:
                    if current:
                        lines.append(current)
                    current = word
            if current:
                lines.append(current)
            return lines

        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # Layout constants (formal document spacing)
        margin = 48
        header_h = 54
        footer_h = 26
        content_top = height - margin - header_h
        content_bottom = margin + footer_h
        max_width = width - (2 * margin)
        # Monochrome palette for professional/legal docs
        ink = colors.HexColor("#111111")
        muted = colors.HexColor("#666666")
        line = colors.HexColor("#CFCFCF")
        y = content_top
        header_drawn = False
        def draw_page_frame():
            # Minimal footer only (no decorative separators)
            c.setFillColor(muted)
            c.setFont("Helvetica", 8.5)
            c.drawString(margin, margin + 10, f"{employer_info['name']} - Employment Contract")
            c.drawRightString(width - margin, margin + 10, f"Page {c.getPageNumber()}")
        def draw_title_block():
            nonlocal y
            c.setFillColor(ink)
            c.setFont("Helvetica-Bold", 15)
            c.drawString(margin, height - margin - 24, title)
            # Generation metadata
            c.setFillColor(muted)
            c.setFont("Helvetica", 9)
            c.drawString(margin, height - margin - 39, f"Generated on {timezone.now().strftime('%Y-%m-%d %H:%M')}")
            y = content_top - 8
        def draw_header_block():
            """Render formal party metadata rows."""
            nonlocal y, header_drawn
            if header_drawn:
                return
            c.setFillColor(ink)
            c.setFont("Helvetica-Bold", 10.2)
            c.drawString(margin, y, "Contract Parties")
            y -= 16
            c.setFont("Helvetica", 9.6)
            rows = [
                ("Prepared by", employer_info["name"] or "-"),
                ("Prepared for", employee_info["name"] or "-"),
                ("Employer date", employer_info["date"] or "-"),
                ("Employee date", employee_info["date"] or "-"),
            ]
            label_width = 98
            row_h = 14
            for label, value in rows:
                ensure_space(20)
                c.setFillColor(muted)
                c.drawString(margin, y, f"{label}:")
                c.setFillColor(ink)
                wrapped_value = wrap_lines(str(value), "Helvetica", 9.6, max_width - label_width, c)
                c.drawString(margin + label_width, y, wrapped_value[0] if wrapped_value else "-")
                y -= row_h
                for continuation in wrapped_value[1:2]:
                    c.drawString(margin + label_width, y, continuation)
                    y -= row_h
            y -= 8
            header_drawn = True

        def new_page():
            c.showPage()
            draw_page_frame()
            draw_title_block()

        def ensure_space(min_needed: float):
            nonlocal y
            if y - min_needed < content_bottom:
                new_page()
                y = content_top - 8

        # Start first page
        draw_page_frame()
        draw_title_block()
        ensure_space(90)
        draw_header_block()

        # Body rendering
        c.setFillColor(ink)
        c.setFont("Helvetica", 10)

        for para in body.split("\n\n"):
            # Flatten paragraph into wrapped lines
            para_lines = []
            for raw_line in para.split("\n"):
                raw_line = raw_line.strip()
                if not raw_line:
                    para_lines.append("")
                    continue
                para_lines.extend(wrap_lines(raw_line, "Helvetica", 10, max_width, c))

            # Render lines
            for line_text in para_lines:
                lower_text = line_text.strip().lower()

                # Skip legacy signature placeholders from templates and header/address lines
                if (
                    lower_text.startswith("12. signatures")
                    or lower_text.startswith("for ")
                    or lower_text.startswith("employee (")
                    or _should_skip_body_line(line_text, employer_info, employee_info)
                ):
                    continue

                # Spacing between empty lines
                if not line_text.strip():
                    y -= 8
                    continue

                ensure_space(24)

                text = line_text.strip()

                # Headings: "1. TERM" style OR all-caps
                is_heading = (text[:1].isdigit() and ". " in text[:6]) or (text.isupper() and len(text) > 3)
                if is_heading:
                    # Section heading styling (no rules/lines)
                    c.setFillColor(ink)
                    c.setFont("Helvetica-Bold", 11.2)
                    c.drawString(margin, y, text)
                    y -= 18
                    c.setFillColor(ink)
                    c.setFont("Helvetica", 10)
                    continue

                # Normal text
                c.setFillColor(ink)
                # Bold important labels
                bold_line = text.lower().startswith("base salary") or text.lower().startswith("allowances") or text.lower().startswith("deductions")
                c.setFont("Helvetica-Bold" if bold_line else "Helvetica", 10)
                c.drawString(margin, y, text)
                y -= 16

            y -= 14  # paragraph spacing

        # ---------------------------------------------------------------------
        # Signature block (two columns: employer left, employee right)
        block_height = 140
        ensure_space(block_height + 24)

        y_sig = y - 6
        if y_sig - block_height < content_bottom:
            new_page()
            y_sig = content_top - 10

        c.setFillColor(ink)
        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin, y_sig, "Signatures")
        y_sig -= 18

        col_gap = 28
        col_width = (width - (2 * margin) - col_gap) / 2
        left_x = margin
        right_x = margin + col_width + col_gap

        def draw_signature_party(info, x, y_top, role_label):
            row_h = 14
            c.setFillColor(ink)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(x, y_top, role_label)

            c.setFont("Helvetica", 9.5)
            c.drawString(x, y_top - row_h, f"Name: {info['name'] or '-'}")
            c.setFillColor(muted)
            c.drawString(x, y_top - (2 * row_h), f"Email: {info['email'] or '-'}")
            c.drawString(x, y_top - (3 * row_h), f"Phone: {info['phone'] or '-'}")
            c.drawString(x, y_top - (4 * row_h), f"Date: {info['date'] or '-'}")

            c.setFillColor(ink)
            c.setFont("Helvetica", 9.5)
            sig_label_y = y_top - (5 * row_h) - 4
            c.drawString(x, sig_label_y, "Signature:")

            if info["sig_path"]:
                try:
                    c.drawImage(
                        info["sig_path"],
                        x + 62,
                        sig_label_y - 2,
                        width=max(80, col_width - 72),
                        height=24,
                        preserveAspectRatio=True,
                        mask="auto",
                    )
                except Exception:
                    c.setFillColor(muted)
                    c.drawString(x + 62, sig_label_y + 8, "Signature on file")
            else:
                c.setFillColor(muted)
                c.drawString(x + 62, sig_label_y + 8, "Signature missing")

        draw_signature_party(employer_info, left_x, y_sig, "Employer")
        draw_signature_party(employee_info, right_x, y_sig, "Employee")

        c.save()

        buffer.seek(0)
        return ContentFile(buffer.getvalue())

    # -------------------------------------------------------------------------
    # 3) Load and Process DOCX
    # -------------------------------------------------------------------------
    db_alias = contract._state.db or "default"
    has_file = bool(getattr(template, "file", None)) and bool(getattr(template.file, "name", None))
    ext = (template.file.name or "").lower() if has_file else ""

    if not has_file:
        try:
            from contracts.management.commands.seed_contract_template import TYPE_BODIES, BASE_BODY
            body_text = _clean_body(template.body_override or TYPE_BODIES.get(contract.contract_type, BASE_BODY))
        except Exception:
            body_text = _clean_body(template.body_override or "")
        content_file = render_pdf_from_body(f"{contract.contract_type} CONTRACT", body_text or "Contract")
        file_name = f"Contract_{contract.contract_id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    elif ext.endswith(".pdf"):
        # Always render a fresh PDF with merged context (do not reuse raw template bytes).
        try:
            from contracts.management.commands.seed_contract_template import TYPE_BODIES, BASE_BODY

            body_text = _clean_body(template.body_override or TYPE_BODIES.get(contract.contract_type, BASE_BODY))
        except Exception:
            body_text = template.body_override or ""

        content_file = render_pdf_from_body(f"{contract.contract_type} CONTRACT", body_text or "Contract")
        file_name = f"Contract_{contract.contract_id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.pdf"

    else:
        # DOCX path with python-docx
        try:
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Inches

            template.file.open("rb")
            doc = Document(template.file)
            template.file.close()

            # Helper to replace text in runs (preserves formatting)
            def replace_text_in_paragraph(paragraph, key, value):
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))

            # Replace placeholders and strip legacy signature/header lines
            for paragraph in doc.paragraphs:
                lower_text = paragraph.text.strip().lower()

                if (
                    lower_text.startswith("12. signatures")
                    or lower_text.startswith("for ")
                    or lower_text.startswith("employee (")
                    or _should_skip_body_line(paragraph.text, employer_info, employee_info)
                ):
                    paragraph.text = ""
                    continue

                for key, value in context.items():
                    if key in paragraph.text:
                        replace_text_in_paragraph(paragraph, key, value)

            # Insert canonical parties metadata at the top of the document
            if doc.paragraphs:
                anchor = doc.paragraphs[0]
                anchor.insert_paragraph_before("")
                anchor.insert_paragraph_before(f"Employee date: {employee_info['date']}")
                anchor.insert_paragraph_before(f"Employer date: {employer_info['date']}")
                anchor.insert_paragraph_before(f"Prepared for: {employee_info['name']}")
                anchor.insert_paragraph_before(f"Prepared by: {employer_info['name']}")
                party_title = anchor.insert_paragraph_before("Contract Parties")
                if party_title.runs:
                    party_title.runs[0].bold = True

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in context.items():
                                if key in paragraph.text:
                                    replace_text_in_paragraph(paragraph, key, value)

            # Add signature section in two columns (Employer left, Employee right)
            doc.add_paragraph("")
            sig_title = doc.add_paragraph("Signatures")
            if sig_title.runs:
                sig_title.runs[0].bold = True

            sig_table = doc.add_table(rows=1, cols=2)

            # Remove table borders to keep the section clean while preserving left/right layout.
            tbl = sig_table._tbl
            tbl_pr = tbl.tblPr
            borders = OxmlElement("w:tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                edge_tag = OxmlElement(f"w:{edge}")
                edge_tag.set(qn("w:val"), "nil")
                borders.append(edge_tag)
            tbl_pr.append(borders)

            def fill_signature_cell(cell, role_label, info):
                cell.text = ""
                role_para = cell.add_paragraph(role_label)
                if role_para.runs:
                    role_para.runs[0].bold = True
                cell.add_paragraph(f"Name: {info['name'] or '-'}")
                cell.add_paragraph(f"Email: {info['email'] or '-'}")
                cell.add_paragraph(f"Phone: {info['phone'] or '-'}")
                cell.add_paragraph(f"Date: {info['date'] or '-'}")

                sig_para = cell.add_paragraph("Signature:")
                if info["sig_path"]:
                    try:
                        sig_para.add_run().add_picture(info["sig_path"], width=Inches(2.0))
                    except Exception:
                        sig_para.add_run(" [Signature on file]")
                else:
                    sig_para.add_run(" Missing signature")

            fill_signature_cell(sig_table.cell(0, 0), "Employer", employer_info)
            fill_signature_cell(sig_table.cell(0, 1), "Employee", employee_info)

            docx_io = BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)

            content_file = ContentFile(docx_io.getvalue())
            file_name = f"Contract_{contract.contract_id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.docx"

        except ImportError:
            # Fallback if python-docx missing: render a simple PDF from body
            from contracts.management.commands.seed_contract_template import TYPE_BODIES, BASE_BODY

            body_text = _clean_body(TYPE_BODIES.get(contract.contract_type, BASE_BODY))
            content_file = render_pdf_from_body(f"{contract.contract_type} CONTRACT", body_text)
            file_name = f"Contract_{contract.contract_id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.pdf"

        except Exception as e:
            raise ValueError(f"Could not open template file: {e}")

    # -------------------------------------------------------------------------
    # 4) Replace any previous generated docs (and their files) to avoid stale files
    # -------------------------------------------------------------------------
    old_docs = ContractDocument.objects.using(db_alias).filter(contract=contract)
    for old in old_docs:
        if old.file:
            old.file.delete(save=False)
        old.delete()

    doc_obj = ContractDocument(
        contract=contract,
        generated_from=template,
        name=file_name,
    )
    doc_obj.file.save(file_name, content_file, save=False)
    doc_obj.save(using=db_alias)
    return doc_obj

